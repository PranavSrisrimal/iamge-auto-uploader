from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import imghdr
import os, os.path

file_list = [name for name in os.listdir('.') if os.path.isfile(name) and imghdr.what(name)]

def paragraph_format_run(cell):
    paragraph = cell.paragraphs[0]
    paragraph = cell.add_paragraph()
    format = paragraph.paragraph_format
    run = paragraph.add_run()
    
    format.space_before = Pt(1)
    format.space_after = Pt(10)
    format.line_spacing = 1.0
    format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return paragraph, format, run

def main():
    document = Document()
    
    sections = document.sections
    section = sections[0]
    
    for section in sections:
    # change orientation to landscape
        section.orientation = WD_ORIENT.LANDSCAPE

        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
    
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    table = document.add_table(rows=2, cols=2)
    table.allow_autofit = False
    
    cells = table.rows[0].cells
    
    #Insert IDs
    file_length = len(file_list)
    k=0
    for i in range(len(file_list)):
        for j in range(2):
            if k < file_length:
                pic_path = file_list[k]
                k += 1
                cell = cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                cell_p, cell_f, cell_r = paragraph_format_run(cell)

                cell_r.add_picture(pic_path)

    doc_path = "IDs.docx"
    document.save(doc_path)

if __name__ == "__main__":    
    main()
